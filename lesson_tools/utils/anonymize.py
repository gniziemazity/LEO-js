import os
import sys
import csv
import re
import shutil
import zipfile
import tempfile
from pathlib import Path

from .similarity_measures import open_csv_encoded

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

_INVALID_FS_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def _safe_folder_name(name):
    cleaned = _INVALID_FS_CHARS.sub("", name).strip().rstrip(".")
    return cleaned


def load_excluded_student_ids(csv_path):
    if not os.path.isfile(csv_path):
        return set()
    excluded = set()
    has_include = False
    for enc in ('utf-8-sig', 'utf-8', 'latin-1', 'cp1252'):
        excluded.clear()
        has_include = False
        try:
            with open(csv_path, encoding=enc, newline='') as fh:
                reader = csv.DictReader(fh, delimiter=';')
                if not reader.fieldnames or 'Include' not in reader.fieldnames:
                    return set()
                has_include = True
                for row in reader:
                    sid = (row.get('Student ID') or '').strip()
                    inc = (row.get('Include') or '').strip()
                    if sid and inc.upper() != 'OK':
                        excluded.add(sid)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception:
            return set()
    return excluded if has_include else set()


def load_students(csv_path):
    students = {}

    def _row(row):
        sid = row["Student ID"].strip()
        name = row["Student Name"].strip()
        number = row["Student Number"].strip()
        alter = (row.get("Alter Ego") or "").strip()
        include_raw = row.get("Include")
        if include_raw is None:
            included = True
        else:
            included = include_raw.strip().upper() in ("OK", "AI")
        students[name] = {
            "id": sid,
            "name": name,
            "number": number,
            "alter_ego": alter,
            "included": included,
        }

    try:
        open_csv_encoded(csv_path, _row, delimiter=";", reset_fn=students.clear)
    except KeyError as e:
        print(f"ERROR: Missing expected column {e} in students.csv")
        print("  Expected columns: Student ID;Student Name;Student Number")
        sys.exit(1)

    if not students:
        print("ERROR: Could not read students.csv with any encoding.")
        sys.exit(1)

    return students

def match_folder_to_student(folder_name, students):
    if folder_name in students:
        return students[folder_name]

    folder_lower = folder_name.lower()
    for name, data in students.items():
        if name.lower() == folder_lower:
            return data

    for name, data in students.items():
        name_lower = name.lower()
        if folder_lower in name_lower or name_lower in folder_lower:
            return data

    folder_parts = set(folder_lower.split())
    for name, data in students.items():
        name_parts = set(name.lower().split())
        overlap = folder_parts & name_parts
        if len(overlap) >= max(1, min(len(folder_parts), len(name_parts)) - 1):
            if len(overlap) >= 1 and len(folder_parts) <= 4:
                return data

    return None

def get_name_patterns(name):
    patterns = []
    skipped_short = []
    seen = set()

    full = name.strip()
    if full:
        patterns.append(re.compile(re.escape(full), re.IGNORECASE))
        seen.add(full.lower())

    for part in name.split():
        key = part.lower()
        if key in seen:
            continue
        if len(part) >= 4:
            patterns.append(re.compile(re.escape(part), re.IGNORECASE))
            seen.add(key)
        elif len(part) == 3:
            patterns.append(re.compile(r"\b" + re.escape(part) + r"\b", re.IGNORECASE))
            seen.add(key)
        else:
            skipped_short.append(part)

    return patterns, skipped_short

def anonymize_text(text, student_data, all_student_numbers):
    remarks = []
    number = student_data["number"]

    text = text.replace(number, "123456")

    patterns, skipped_short = get_name_patterns(student_data["name"])
    for pattern in patterns:
        text = pattern.sub("XXX", text)
    if skipped_short:
        remarks.append(
            "Short name part(s) not auto-redacted (verify manually): "
            + ", ".join(skipped_short)
        )

    found_numbers = re.findall(r"\b(\d{7})\b", text)
    for found in found_numbers:
        if found == "123456":
            continue
        if found in all_student_numbers:
            remarks.append(f"Contains another student number: {found}")
        else:
            remarks.append(f"Contains unknown digit sequence: {found}")

    return text, remarks

def anonymize_filename(filename, student_data):
    new_name = filename

    new_name = new_name.replace(student_data["number"], "123456")

    patterns, _ = get_name_patterns(student_data["name"])
    for pattern in patterns:
        new_name = pattern.sub("XXX", new_name)

    return new_name

def read_text_file(filepath):
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            with open(filepath, "r", encoding=enc) as f:
                content = f.read()
            return content, enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return None, None

def process_text_file(src_path, dst_path, student_data, all_student_numbers):
    remarks = []
    content, encoding = read_text_file(src_path)

    if content is not None:
        content, file_remarks = anonymize_text(content, student_data, all_student_numbers)
        remarks.extend(file_remarks)
        with open(dst_path, "w", encoding=encoding) as f:
            f.write(content)
    else:
        shutil.copy2(src_path, dst_path)
        remarks.append(f"Could not decode: {os.path.basename(src_path)}")

    return remarks

_WORD_DOC_CT = (
    b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
)


def _open_word_document(src_path):
    try:
        return Document(src_path), None
    except ValueError:
        pass

    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    with zipfile.ZipFile(src_path) as zin, \
            zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = re.sub(
                    rb'(<Override PartName="/word/document\.xml" ContentType=")[^"]*(")',
                    rb"\g<1>" + _WORD_DOC_CT + rb"\g<2>",
                    data,
                )
            zout.writestr(item, data)
    return Document(tmp_path), tmp_path


def process_docx_file(src_path, dst_path, student_data, all_student_numbers):
    remarks = []
    if not HAS_DOCX:
        shutil.copy2(src_path, dst_path)
        remarks.append(
            f"python-docx not installed, could not anonymize: {os.path.basename(src_path)}"
        )
        return remarks

    tmp_path = None
    try:
        doc, tmp_path = _open_word_document(src_path)
        for para in doc.paragraphs:
            for run in para.runs:
                new_text, run_remarks = anonymize_text(
                    run.text, student_data, all_student_numbers
                )
                run.text = new_text
                remarks.extend(run_remarks)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            new_text, run_remarks = anonymize_text(
                                run.text, student_data, all_student_numbers
                            )
                            run.text = new_text
                            remarks.extend(run_remarks)

        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf is not None:
                    for para in hf.paragraphs:
                        for run in para.runs:
                            new_text, run_remarks = anonymize_text(
                                run.text, student_data, all_student_numbers
                            )
                            run.text = new_text
                            remarks.extend(run_remarks)

        cp = doc.core_properties
        had_author = bool((cp.author or "").strip() or (cp.last_modified_by or "").strip())
        for field in ("author", "last_modified_by", "title", "subject",
                      "keywords", "comments", "category", "content_status"):
            try:
                setattr(cp, field, "")
            except (ValueError, AttributeError):
                pass
        if had_author:
            remarks.append(f"DOCX: cleared author metadata ({os.path.basename(src_path)})")

        doc.save(dst_path)
    except Exception as e:
        shutil.copy2(src_path, dst_path)
        remarks.append(f"DOCX error ({os.path.basename(src_path)}): {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    return remarks

def process_pdf_file(src_path, dst_path, student_data, all_student_numbers):
    remarks = []
    if not HAS_PDF:
        shutil.copy2(src_path, dst_path)
        remarks.append(
            f"PyMuPDF not installed, could not anonymize: {os.path.basename(src_path)}"
        )
        return remarks

    try:
        doc = fitz.open(src_path)
        number = student_data["number"]
        name = student_data["name"]

        parts = name.split()
        name_terms = []
        seen_terms = set()
        term_candidates = [name]
        if len(parts) > 1:
            term_candidates.append(" ".join(reversed(parts)))
        term_candidates.extend(p for p in parts if len(p) >= 4)
        for term in term_candidates:
            term = term.strip()
            if term and term.lower() not in seen_terms:
                name_terms.append(term)
                seen_terms.add(term.lower())

        skipped_short = [p for p in parts if len(p) < 4]
        if skipped_short:
            remarks.append(
                "Short name part(s) not auto-redacted in PDF (verify manually): "
                + ", ".join(skipped_short)
            )

        for page_num, page in enumerate(doc, 1):
            num_instances = page.search_for(number)
            if num_instances:
                for inst in num_instances:
                    page.add_redact_annot(inst, text="123456")
                remarks.append(
                    f"PDF p.{page_num}: redacted student number ({len(num_instances)}x)"
                )

            name_hits = 0
            for term in name_terms:
                for inst in page.search_for(term):
                    page.add_redact_annot(inst, text="XXX")
                    name_hits += 1
            if name_hits:
                remarks.append(
                    f"PDF p.{page_num}: redacted student name ({name_hits}x)"
                )

            text = page.get_text()
            found_numbers = re.findall(r"\b(\d{7})\b", text)
            for found in found_numbers:
                if found == number or found == "123456":
                    continue
                if found in all_student_numbers:
                    remarks.append(
                        f"PDF p.{page_num}: contains another student number: {found}"
                    )

            page.apply_redactions()

        had_author = bool((doc.metadata or {}).get("author", "").strip())
        doc.set_metadata({})
        try:
            doc.del_xml_metadata()
        except AttributeError:
            pass
        if had_author:
            remarks.append(f"PDF: cleared author metadata ({os.path.basename(src_path)})")

        doc.save(dst_path, garbage=4, deflate=True, clean=True)
        doc.close()
    except Exception as e:
        shutil.copy2(src_path, dst_path)
        remarks.append(f"PDF error ({os.path.basename(src_path)}): {e}")

    return remarks

TEXT_EXTENSIONS = {
    ".html", ".htm", ".css", ".js", ".ts", ".json", ".xml", ".svg",
    ".txt", ".md", ".csv", ".yml", ".yaml", ".jsx", ".tsx", ".php",
    ".py", ".java", ".c", ".cpp", ".h", ".rb", ".sh", ".bat",
}
DOCX_EXTENSIONS = {".docx", ".dotx", ".docm", ".dotm"}
PDF_EXTENSIONS = {".pdf"}

VALID_STUDENT_EXTENSIONS = {'.html', '.htm', '.css', '.js', '.py', '.docx', '.pdf'}

def process_file(src_path, dst_path, student_data, all_student_numbers):
    ext = os.path.splitext(src_path)[1].lower()

    remarks = []
    filename = os.path.basename(src_path)
    if ext not in VALID_STUDENT_EXTENSIONS and filename != 'tokens.txt':
        remarks.append(f"Unexpected file type: {filename}")

    if ext in TEXT_EXTENSIONS:
        remarks.extend(process_text_file(src_path, dst_path, student_data, all_student_numbers))
    elif ext in DOCX_EXTENSIONS:
        remarks.extend(process_docx_file(src_path, dst_path, student_data, all_student_numbers))
    elif ext in PDF_EXTENSIONS:
        remarks.extend(process_pdf_file(src_path, dst_path, student_data, all_student_numbers))
    else:
        shutil.copy2(src_path, dst_path)

    return remarks

def main():
    if len(sys.argv) < 2:
        print("Usage: anonymize.py <project_dir>")
        sys.exit(1)
    project_dir = sys.argv[1]

    students_csv = os.path.join(project_dir, "..", "..", "students.csv")
    students_dir = os.path.join(project_dir, "students")
    anon_names_dir = os.path.join(project_dir, "anon_names")
    anon_ids_dir = os.path.join(project_dir, "anon_ids")
    remarks_csv = os.path.join(project_dir, "remarks.csv")

    if not os.path.exists(students_csv):
        print("ERROR: students.csv not found! This file is required.")
        print(f"  Expected at: {students_csv}")
        sys.exit(1)

    if not os.path.exists(students_dir):
        print("ERROR: students/ folder not found!")
        print(f"  Expected at: {students_dir}")
        sys.exit(1)

    students = load_students(students_csv)
    all_student_numbers = {s["number"] for s in students.values()}

    print(f"Loaded {len(students)} students from students.csv")
    print("  Columns found: Student ID, Student Name, Student Number")

    if not HAS_DOCX:
        print("  Note: python-docx not installed (pip install python-docx). "
              "DOCX files will be copied as-is.")
    if not HAS_PDF:
        print("  Note: PyMuPDF not installed (pip install pymupdf). "
              "PDF files will be copied as-is.")

    for d in [anon_names_dir, anon_ids_dir]:
        if os.path.exists(d):
            shutil.rmtree(d)
        os.makedirs(d)

    student_folders = [
        f
        for f in os.listdir(students_dir)
        if os.path.isdir(os.path.join(students_dir, f))
    ]

    print(f"Found {len(student_folders)} student folders in students/\n")

    processed_remarks = {}
    matched = 0
    unmatched = 0

    use_alter_ego = os.environ.get("STUDENT_ANALYTICS_USE_ALTER_EGO") == "1"

    skipped_excluded = 0

    for folder_name in sorted(student_folders):
        student = match_folder_to_student(folder_name, students)

        if student is None:
            print(f"  WARNING: No CSV match for folder '{folder_name}' -- skipping")
            unmatched += 1
            continue

        if not student.get("included", True):
            print(f"  Skipping '{folder_name}' (Include != OK in students.csv)")
            skipped_excluded += 1
            continue

        matched += 1
        src_folder = os.path.join(students_dir, folder_name)

        names_folder = folder_name
        if use_alter_ego and student.get("alter_ego"):
            names_folder = _safe_folder_name(student["alter_ego"]) or folder_name
        dest_names = os.path.join(anon_names_dir, names_folder)
        dest_ids = os.path.join(anon_ids_dir, student["id"])
        os.makedirs(dest_names, exist_ok=True)
        os.makedirs(dest_ids, exist_ok=True)

        folder_remarks = []
        number_locations = []

        number = student["number"]
        number_pattern = re.compile(r'(?<!\d)' + re.escape(number) + r'(?!\d)')

        for root, dirs, files in os.walk(src_folder):
            rel_path = os.path.relpath(root, src_folder)

            for filename in files:
                if filename.lower() == "onlinetext.txt":
                    folder_remarks.append("Removed onlinetext.txt (identity-revealing links)")
                    continue

                src_file = os.path.join(root, filename)
                ext = os.path.splitext(src_file)[1].lower()

                if number in filename:
                    display = filename if rel_path == "." else os.path.join(rel_path, filename)
                    number_locations.append(f"filename: {display}")

                if ext in TEXT_EXTENSIONS:
                    content, _ = read_text_file(src_file)
                    if content and number_pattern.search(content):
                        display = filename if rel_path == "." else os.path.join(rel_path, filename)
                        number_locations.append(f"in {display}")

                elif ext in DOCX_EXTENSIONS and HAS_DOCX:
                    scan_tmp = None
                    try:
                        doc_obj, scan_tmp = _open_word_document(src_file)
                        full_text = "\n".join(
                            run.text
                            for para in doc_obj.paragraphs
                            for run in para.runs
                        )
                        if number_pattern.search(full_text):
                            display = filename if rel_path == "." else os.path.join(rel_path, filename)
                            number_locations.append(f"in {display}")
                    except Exception:
                        pass
                    finally:
                        if scan_tmp and os.path.exists(scan_tmp):
                            try:
                                os.remove(scan_tmp)
                            except OSError:
                                pass

                elif ext in PDF_EXTENSIONS and HAS_PDF:
                    try:
                        import fitz
                        pdf_doc = fitz.open(src_file)
                        for page in pdf_doc:
                            if number_pattern.search(page.get_text()):
                                display = filename if rel_path == "." else os.path.join(rel_path, filename)
                                number_locations.append(f"in {display}")
                                break
                        pdf_doc.close()
                    except Exception:
                        pass

                anon_fname = anonymize_filename(filename, student)

                if rel_path == ".":
                    dst_names_path = os.path.join(dest_names, anon_fname)
                    dst_ids_path = os.path.join(dest_ids, anon_fname)
                else:
                    sub_names = os.path.join(dest_names, rel_path)
                    sub_ids = os.path.join(dest_ids, rel_path)
                    os.makedirs(sub_names, exist_ok=True)
                    os.makedirs(sub_ids, exist_ok=True)
                    dst_names_path = os.path.join(sub_names, anon_fname)
                    dst_ids_path = os.path.join(sub_ids, anon_fname)

                file_remarks = process_file(
                    src_file, dst_names_path, student, all_student_numbers
                )
                folder_remarks.extend(file_remarks)

                process_file(src_file, dst_ids_path, student, all_student_numbers)

        if not number_locations:
            folder_remarks.insert(0, "Student number not found")

        remark_text = "; ".join(dict.fromkeys(folder_remarks))
        processed_remarks[student["number"]] = remark_text

        tag = f" >> {len(folder_remarks)} remark(s)" if folder_remarks else ""
        print(f"  [{student['id']:>2}] {folder_name}{tag}")

    with open(remarks_csv, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow(["Student Number", "Remarks"])
        for number, text in processed_remarks.items():
            writer.writerow([number, text])

    print(f"\n{'=' * 50}")
    print("Done!")
    print(f"  Matched:   {matched} student folders")
    print(f"  Unmatched: {unmatched} student folders")
    if skipped_excluded:
        print(f"  Skipped:   {skipped_excluded} student folder(s) (Include != OK)")
    print(f"  anon_names/ -- folders named by student name, content anonymized")
    print(f"  anon_ids/   -- folders named by student ID, content anonymized")
    print(f"  remarks.csv -- {len(processed_remarks)} entries written")
    print(f"{'=' * 50}")

if __name__ == "__main__":
    main()